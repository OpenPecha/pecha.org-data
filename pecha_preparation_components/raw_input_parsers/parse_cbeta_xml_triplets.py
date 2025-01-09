import re
from pathlib import Path
from collections import defaultdict
from docx import Document  # from bayoo_docx


def find_triplet(files):
    tri = defaultdict(dict)
    for f in files:
        stem = f.stem
        stem = stem if '.' not in stem else stem.split('.')[0]
        if '.bo.zh' in f.name or '.bo.cn' in f.name or '.Tib.Chn' in f.name or '.BO.CN' in f.name:
            tri[stem]['bo_zh'] = f
        elif '.zh' in f.name or '.cn' in f.name or '.Chn' in f.name or '.CN' in f.name:
            tri[stem]['zh'] = f
        elif '.bo' in f.name or '.BO' in f.name or '.Tib' in f.name:
            tri[stem]['bo'] = f
        else:
            print('there is a problem')
            print(f)
    return tri


def parse_triplets(in_folder):
    # requires xml files triplets each in a sub_folder.  no more than 1 level
    modern_folder_name = "Modern Chinese"
    total = {}
    for folder in Path(in_folder).glob("*"):
        files = list(folder.glob("*.xml"))
        sub_dirs = [sub for sub in folder.glob("*") if sub.is_dir() and sub.name == modern_folder_name]
        tri = {'trad': find_triplet(files), 'simp': None}
        if sub_dirs:
            files = list(sub_dirs[0].glob("*.xml"))
            tri['simp'] = find_triplet(files)
        total[folder.name] = tri

    # check every triplet has 3 files and extract incomplete triplets
    incomplete = {}
    for folder, tr_sm in total.items():
        for h, tri in tr_sm.items():
            if not tri:
                continue
            keys = list(tri.keys())
            for k in keys:
                v = tri[k]
                if len(v) != 3:
                    incomplete[folder] = {k: v}
                    del total[folder][h][k]
    return total, incomplete

def parse_lang(in_file):
    dump = in_file.read_text()
    lines = dump.split('\n')
    out = {}
    for line in lines:
        if '"/>' in line:
            a = re.findall(r's id=\"([0-9\:]+)\"\/>', line)
            a.append('')
        else:
            a = re.findall(r's id=\"([0-9\:]+)\">([^<]+)<\/s>', line)
            a = a[0] if a else []
        if a:
            out[a[0]] = a[1]
    return out

def parse_table(in_file):
    dump = in_file.read_text()
    a = re.findall(r"type='([0-9\-]+)' xtargets='([0-9\:]*?)\;([0-9\:]*?)' status", dump)
    out = []
    for i, j, k in a:
        out.append({'type': i, 'zh': j, 'bo': k})
    return out

def align_triplet(triplet):
    # get file name
    name = triplet['bo'].name.split('.')[0]
    print('\t', name)

    # parse xml files
    lang1 = parse_lang(triplet['bo'])
    lang2 = parse_lang(triplet['zh'])
    table = parse_table(triplet['bo_zh'])

    # align
    aligned = []
    for t in table:
        bo, zh = t['bo'], t['zh']
        entry = ['', '']
        if bo:
            try:
                entry[0] = lang1[bo]
            except:
                print('\t\tbo: ', bo)
        if zh:
            try:
                entry[1] = lang2[zh]
            except:
                print('\t\tzh: ', zh)
        aligned.append(entry)
    return name, aligned

def write_documents(folder, filename, sim_trad, aligned):
    # export to docx and parse footnotes
    if not folder.exists():
        folder.mkdir(parents=True, exist_ok=True)
    doc_bo = Document()
    doc_zh = Document()
    line_num = 1
    for bo, zh in aligned:
        # add Tibetan
        doc_bo.add_paragraph(f'{line_num}. {bo}')

        # add Chinese text parsing footnotes as required
        if '^' in zh:
            par = doc_zh.add_paragraph()
            # extract notes from the string
            footnotes = re.findall(r'\^[0-9]+\[([^\]]+)\]', zh)
            # replace notes by * and split on * to know where to add footnotes
            parts = re.sub(r'\^[0-9]+\[[^\]]+?\]', 'ᚸ', zh)
            parts = re.split(r'(ᚸ)', parts)
            # adding text with footnotes at correct places
            start_added = False
            f_num = 0
            for p in parts:
                if p != 'ᚸ':
                    if not start_added:
                        par.add_run(f'{line_num}. {p}')
                        start_added = True
                    else:
                        par.add_run(p)
                else:
                    try:
                        par.add_footnote(footnotes[f_num])
                        f_num += 1
                    except:
                        print("unable to parse note:", zh)
        else:
            doc_zh.add_paragraph(f'{line_num}. {zh}')
        line_num += 1

    if sim_trad == 'simp':
        st = 'simplified_'
    else:
        st = ''
    doc_bo.save(folder / f'{filename}_{st}bo.docx')
    doc_zh.save(folder / f'{filename}_{st}zh.docx')

def parse_cbeta_xml_triplets(in_folder, out_folder):
    out_folder = out_folder / 'Gold Standard'

    triplets, incomplete = parse_triplets(in_folder)
    for work, parts in triplets.items():
        print(work)
        cur_out_folder = out_folder / work
        for h, p in parts.items():
            if not p:
                continue
            out = []
            sorted_parts = [p[s] for s in sorted(p.keys())]
            for tri in sorted_parts:
                _, aligned = align_triplet(tri)
                out.extend(aligned)
            write_documents(cur_out_folder, work, h, out)